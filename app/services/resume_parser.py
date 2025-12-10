import io
import tempfile
from pathlib import Path
from typing import Optional
import logging

import pdfplumber
from docx import Document

from ..schemas import ExtractedData
from .text_extractors import TextExtractor
from .llm_extractor import LLMExtractor

logger = logging.getLogger(__name__)

class ResumeParser:
    """Service for parsing resume files and extracting structured data"""
    
    def __init__(self):
        self.text_extractor = TextExtractor()
        self.llm_extractor = LLMExtractor()
    
    def parse_resume(self, filename: str, file_content: bytes) -> ExtractedData:
        """
        Parse a resume file and extract structured data
        
        Args:
            filename: Name of the file
            file_content: Binary content of the file
            
        Returns:
            ExtractedData object with extracted information
        """
        try:
            # Extract text based on file extension
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_content)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_content)
            else:
                logger.error(f"Unsupported file type: {file_ext}")
                return ExtractedData()  # Return empty data
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from {filename}")
                return ExtractedData()
            
            logger.info(f"Extracted {len(text)} characters from {filename}")
            
            # Try LLM extraction first (if enabled)
            extracted_data = self.llm_extractor.extract_data(text)
            
            # Fallback to spaCy/regex extraction if LLM fails
            if extracted_data is None:
                logger.info(f"Using spaCy extraction for {filename}")
                extracted_data = self.text_extractor.extract_data(text)
            else:
                logger.info(f"Used LLM extraction for {filename}")
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {str(e)}")
            return ExtractedData()  # Return empty data on error
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file content"""
        try:
            text_parts = []
            
            with io.BytesIO(file_content) as pdf_stream:
                with pdfplumber.open(pdf_stream) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file content"""
        try:
            text_parts = []
            
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
            
            finally:
                # Clean up temporary file
                Path(temp_file_path).unlink(missing_ok=True)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return ""